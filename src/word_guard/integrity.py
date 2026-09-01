# -*- coding: utf-8 -*-
"""
Scope and formatting integrity checks for DOCX.

Text equality alone is not enough for scholarly Word revision: unchanged text
can still be damaged by run formatting, paragraph styles, equations, drawings,
or table mutations.  This module therefore exposes a backwards-compatible
text scope check plus optional structural fingerprints.
"""

from dataclasses import dataclass, field
import hashlib
from lxml import etree
from docx.oxml.ns import qn


@dataclass
class IntegrityCheckResult:
    scope_intact: bool = True
    requested_scope: str = ""
    unchanged_regions: list = field(default_factory=list)
    unexpected_changes: list = field(default_factory=list)
    warnings: list = field(default_factory=list)
    structural_changes: list = field(default_factory=list)


def _get_para_texts(doc):
    return [p.text for p in doc.paragraphs]


def _canonical_xml(element):
    """Stable XML bytes for exact structural comparison."""
    try:
        return etree.tostring(element, method="c14n", with_comments=False)
    except Exception:
        return etree.tostring(element)


def _hash_element(element):
    return hashlib.sha256(_canonical_xml(element)).hexdigest()


def _paragraph_fingerprint(para):
    """
    Fingerprint paragraph style/run formatting and protected inline objects.

    Text is reported separately, so this hash intentionally covers the full
    paragraph XML to catch formatting-only drift.
    """
    return _hash_element(para._p)


def _table_fingerprint(table):
    return _hash_element(table._tbl)


def check_scope_integrity(
    doc_before, doc_after, scope, scope_description="", structural=True
):
    """
    Verify that only the intended paragraph scope was modified.

    ``structural=True`` additionally detects formatting/object drift outside
    scope and changes to tables that were not explicitly included in the scope.
    """
    result = IntegrityCheckResult(requested_scope=scope_description)
    texts_before = _get_para_texts(doc_before)
    texts_after = _get_para_texts(doc_after)

    if len(texts_before) != len(texts_after):
        result.warnings.append(
            f"Paragraph count changed: {len(texts_before)} → {len(texts_after)}"
        )

    scope_indices = set(scope.included_paragraphs)
    total_paras = min(len(texts_before), len(texts_after))

    for i in range(total_paras):
        if i in scope_indices:
            continue
        if texts_before[i] != texts_after[i]:
            result.unexpected_changes.append({
                "type": "text",
                "paragraph_index": i,
                "before": texts_before[i][:100],
                "after": texts_after[i][:100],
            })
        elif structural:
            before_fp = _paragraph_fingerprint(doc_before.paragraphs[i])
            after_fp = _paragraph_fingerprint(doc_after.paragraphs[i])
            if before_fp != after_fp:
                result.structural_changes.append({
                    "type": "paragraph_format_or_object",
                    "paragraph_index": i,
                    "text": texts_before[i][:100],
                })

    scope_changed = sum(
        1 for i in range(total_paras)
        if i in scope_indices and texts_before[i] != texts_after[i]
    )
    scope_unchanged = sum(
        1 for i in range(total_paras)
        if i in scope_indices and texts_before[i] == texts_after[i]
    )

    # Tables are separately indexed by python-docx. Only tables explicitly
    # included by scope are permitted to mutate.
    if structural:
        included_tables = set(getattr(scope, "included_tables", []) or [])
        table_count = min(len(doc_before.tables), len(doc_after.tables))
        if len(doc_before.tables) != len(doc_after.tables):
            result.structural_changes.append({
                "type": "table_count",
                "before": len(doc_before.tables),
                "after": len(doc_after.tables),
            })
        for i in range(table_count):
            if i in included_tables:
                continue
            if _table_fingerprint(doc_before.tables[i]) != _table_fingerprint(doc_after.tables[i]):
                result.structural_changes.append({
                    "type": "table_structure_or_format",
                    "table_index": i,
                })

    if not result.unexpected_changes and not result.structural_changes:
        result.scope_intact = True
        result.unchanged_regions.append(
            "All out-of-scope text and checked structure are unchanged"
        )
    else:
        result.scope_intact = False
        if result.unexpected_changes:
            result.warnings.append(
                f"{len(result.unexpected_changes)} unexpected text changes found outside scope"
            )
        if result.structural_changes:
            result.warnings.append(
                f"{len(result.structural_changes)} unexpected structural/formatting changes found outside scope"
            )

    result.unchanged_regions.append(
        f"Scope: {scope_changed} paragraphs modified, "
        f"{scope_unchanged} paragraphs unchanged within scope"
    )
    return result


def compare_paragraphs(before_paras, after_paras, indices):
    results = []
    for i in indices:
        if i < len(before_paras) and i < len(after_paras):
            before_text = before_paras[i].text
            after_text = after_paras[i].text
            results.append({
                "index": i,
                "changed": before_text != after_text,
                "structure_changed": (
                    _paragraph_fingerprint(before_paras[i])
                    != _paragraph_fingerprint(after_paras[i])
                ),
                "before": before_text[:200],
                "after": after_text[:200],
            })
    return results
