# -*- coding: utf-8 -*-
"""
Word Manuscript Guard - CLI and Tool Interface.

Provides the main entry points for DOCX operations:
- writing_word_scan: structural scan
- writing_word_edit: safe section/paragraph editing
- writing_word_audit: writing audit on extracted text
- writing_word_scope_check: verify scope integrity
"""

import os
import sys
import json
import copy
import shutil
from pathlib import Path

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from word_guard.scanner import scan_docx, print_scan_summary
from word_guard.scope import resolve_scope, find_paragraph_containing
from word_guard.editor import (
    replace_text_in_paragraph,
    replace_entire_paragraph,
    insert_paragraph_after,
    replace_section_text,
)
from word_guard.protected import check_complex_paragraph, get_protected_elements_in_range
from word_guard.manifest import ChangeManifest
from word_guard.integrity import check_scope_integrity
from word_guard.package_guard import validate_docx_package, compare_package_parts
from word_guard.fingerprint import build_baseline_fingerprint
from word_guard.equation_audit import audit_equations


# ═══════════════════════════════════════════════════════════════════════════
# Tool 1: writing_word_scan
# ═══════════════════════════════════════════════════════════════════════════

def writing_word_scan(file_path):
    """
    Perform a structural scan of a DOCX file.

    Returns document profile, heading hierarchy, paragraph info,
    tables, equations, and warnings.

    Args:
        file_path: path to the .docx file

    Returns:
        dict with complete structural analysis
    """
    if not os.path.exists(file_path):
        return {'error': f"File not found: {file_path}"}

    info = scan_docx(file_path)

    result = {
        'file': info.file_name,
        'profile': info.detected_profile,
        'paragraphs': info.paragraph_count,
        'headings': [
            {
                'index': h.index,
                'level': h.heading_level,
                'text': h.text[:120],
                'style': h.style_id,
            }
            for h in info.headings
        ],
        'tables': [
            {
                'index': t.index,
                'rows': t.rows,
                'cols': t.cols,
                'preview': t.text_preview,
            }
            for t in info.tables
        ],
        'objects': {
            'equations': info.total_equations,
            'drawings': info.total_drawings,
            'fields': info.total_fields,
            'hyperlinks': info.total_hyperlinks,
            'bookmarks': info.total_bookmarks,
            'footnotes': info.total_footnotes,
            'comments': info.total_comments,
        },
        'complex_paragraphs': sum(1 for p in info.paragraphs if p.is_complex),
        'sections': info.section_count,
        'warnings': info.warnings,
    }

    return result


# ═══════════════════════════════════════════════════════════════════════════
# Tool 2: writing_word_edit
# ═══════════════════════════════════════════════════════════════════════════

def writing_word_edit(file_path, replacements, scope_config=None,
                      mode='text_only', output_path=None, backup=True):
    """
    Safely edit specific sections of a DOCX file.

    Args:
        file_path: path to the source .docx file
        replacements: list of dicts with 'old' and 'new' keys
        scope_config: dict with scope resolution parameters:
            - start_heading, end_heading: section range
            - heading: single section to edit
            - start_index, end_index: direct paragraph indices
        mode: 'text_only' (default), 'structural', or 'format_normalization'
        output_path: output file path (default: overwrite original)
        backup: whether to create a backup before editing

    Returns:
        dict with edit results and change manifest
    """
    if not os.path.exists(file_path):
        return {'error': f"File not found: {file_path}"}

    # Preserve a package-level preimage for structural integrity comparison.
    import tempfile
    preimage_fd, preimage_path = tempfile.mkstemp(suffix='.docx', prefix='wg_preimage_')
    os.close(preimage_fd)
    shutil.copy2(file_path, preimage_path)

    # Create backup if needed
    if backup and output_path is None:
        backup_path = file_path + '.bak'
        shutil.copy2(file_path, backup_path)

    # Load document
    from docx import Document
    doc = Document(file_path)

    # Resolve scope
    if scope_config:
        scope = resolve_scope(doc.paragraphs, **scope_config)
        if scope.ambiguous:
            return {
                'error': 'Ambiguous scope',
                'warnings': scope.warnings,
                'options': scope.ambiguity_options,
            }
    else:
        # Full document scope
        scope = resolve_scope(doc.paragraphs)

    # Check protected elements in scope
    protected = get_protected_elements_in_range(
        doc, scope.start_para_index, scope.end_para_index + 1
    )

    # Initialize manifest
    manifest = ChangeManifest(
        file_path,
        requested_scope=scope_config or "full document"
    )

    # Apply replacements
    results = []
    for rep in replacements:
        old_text = rep.get('old', '')
        new_text = rep.get('new', '')

        if not old_text:
            continue

        # Search within scope
        found = False
        for idx in scope.included_paragraphs:
            if idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                if old_text in para.text:
                    # Check complexity
                    complexity = check_complex_paragraph(para)

                    if mode == 'text_only' and complexity['is_complex']:
                        # Warn but still attempt run-aware replacement
                        result = replace_text_in_paragraph(para, old_text, new_text, red=True)
                        result['complexity_warning'] = (
                            f"Paragraph at index {idx} is complex "
                            f"({complexity['risk_level']} risk): "
                            f"{complexity['protected_nodes']}"
                        )
                    else:
                        result = replace_text_in_paragraph(para, old_text, new_text, red=True)

                    result['paragraph_index'] = idx
                    results.append(result)

                    # Record in manifest
                    manifest.add_change(
                        change_type='replace_text',
                        target=f"paragraph {idx}",
                        old_text=old_text,
                        new_text=new_text,
                        paragraph_index=idx,
                        method=result.get('method'),
                        success=result.get('success', False),
                        warnings=result.get('warnings', []),
                    )

                    found = True
                    break

        if not found:
            results.append({
                'success': False,
                'warnings': [f"Text not found in scope: '{old_text[:50]}...'"],
            })
            manifest.add_change(
                change_type='replace_text',
                target='not found',
                old_text=old_text,
                new_text=new_text,
                success=False,
                warnings=[f"Text not found in scope"],
            )

    # Record preserved elements
    manifest.add_preserved(f"Chapters outside scope: paragraphs 0-{scope.start_para_index-1}")
    manifest.add_preserved(f"Chapters outside scope: paragraphs {scope.end_para_index+1}+")
    if protected:
        for p in protected:
            manifest.add_preserved(
                f"Complex paragraph at {p['paragraph_index']}: "
                f"{', '.join(p['protected_nodes'])}"
            )

    # Save
    save_path = output_path or file_path
    doc.save(save_path)
    manifest.finalize()

    package_validation = validate_docx_package(save_path).to_dict()
    # In text-only mode, only the main document XML and volatile core metadata may change.
    allowed_parts = {'word/document.xml', 'docProps/core.xml'}
    package_integrity = compare_package_parts(preimage_path, save_path, allowed_parts)
    try:
        os.unlink(preimage_path)
    except OSError:
        pass

    success = package_validation['valid'] and (package_integrity['intact'] or mode != 'text_only')
    return {
        'success': success,
        'output': save_path,
        'changes': len([r for r in results if r.get('success')]),
        'total_replacements': len(replacements),
        'scope': {
            'start': scope.start_para_index,
            'end': scope.end_para_index,
            'heading': scope.start_heading,
        },
        'protected_elements': len(protected),
        'package_validation': package_validation,
        'package_integrity': package_integrity,
        'manifest_summary': manifest.summary(),
    }


# ═══════════════════════════════════════════════════════════════════════════
# Tool 3: writing_word_audit
# ═══════════════════════════════════════════════════════════════════════════

def writing_word_audit(file_path):
    """
    Extract normalized text from a DOCX and check for writing issues.

    Checks:
    - Revision residue patterns
    - AI-style language
    - Scientific notation issues
    - Unit formatting

    Args:
        file_path: path to the .docx file

    Returns:
        dict with extracted text and audit findings
    """
    if not os.path.exists(file_path):
        return {'error': f"File not found: {file_path}"}

    from docx import Document
    doc = Document(file_path)

    # Extract normalized text with source mapping
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            paragraphs.append({
                'index': i,
                'text': para.text,
                'style': para.style.style_id if para.style else None,
            })

    # Basic writing checks
    issues = []

    # Revision residue patterns (for manuscript profile)
    revision_residue_patterns = [
        'in the original formulation',
        'we revised',
        'the reviewer suggested',
        'in response to the reviewer',
        'the revised model',
        'the previous version',
        'as requested',
        'we have modified',
        'we changed',
    ]

    basename = os.path.basename(file_path).lower()
    is_manuscript = not any(kw in basename for kw in ['rebuttal', 'response', 'reply'])

    for p in paragraphs:
        text_lower = p['text'].lower()
        for pattern in revision_residue_patterns:
            if pattern in text_lower:
                if is_manuscript:
                    issues.append({
                        'type': 'revision_residue',
                        'paragraph': p['index'],
                        'text': p['text'][:100],
                        'pattern': pattern,
                        'severity': 'high',
                    })

    # AI-style patterns
    ai_patterns = [
        ('delve', 'high'),
        ('tapestry', 'high'),
        ('multifaceted', 'medium'),
        ('paradigm', 'medium'),
        ('landscape', 'medium'),
        ('crucial', 'low'),
        ('furthermore', 'low'),
        ('moreover', 'low'),
        ('notably', 'low'),
    ]

    for p in paragraphs:
        text_lower = p['text'].lower()
        for pattern, severity in ai_patterns:
            if pattern in text_lower:
                issues.append({
                    'type': 'ai_style',
                    'paragraph': p['index'],
                    'text': p['text'][:100],
                    'pattern': pattern,
                    'severity': severity,
                })

    # Unit formatting issues
    unit_patterns = [
        ('min-1', 'Should be min⁻¹ or min^-1'),
        ('m-3', 'Should be m⁻³ or m^-3'),
        ('s-1', 'Should be s⁻¹ or s^-1'),
        ('kg m-3', 'Should be kg·m⁻³'),
    ]

    for p in paragraphs:
        for pattern, suggestion in unit_patterns:
            if pattern in p['text']:
                issues.append({
                    'type': 'unit_format',
                    'paragraph': p['index'],
                    'text': p['text'][:100],
                    'pattern': pattern,
                    'suggestion': suggestion,
                    'severity': 'medium',
                })

    return {
        'file': os.path.basename(file_path),
        'profile': 'manuscript' if is_manuscript else 'rebuttal',
        'total_paragraphs': len(paragraphs),
        'issues_found': len(issues),
        'issues': issues,
        'text_preview': [p['text'][:80] for p in paragraphs[:10]],
    }


# ═══════════════════════════════════════════════════════════════════════════
# Tool 4: writing_word_scope_check
# ═══════════════════════════════════════════════════════════════════════════

def writing_word_scope_check(file_before, file_after, scope_config=None):
    """
    Verify that editing did not modify content outside the requested scope.

    Args:
        file_before: path to the original .docx file
        file_after: path to the edited .docx file
        scope_config: dict with scope parameters

    Returns:
        dict with integrity check results
    """
    from docx import Document

    if not os.path.exists(file_before):
        return {'error': f"File not found: {file_before}"}
    if not os.path.exists(file_after):
        return {'error': f"File not found: {file_after}"}

    doc_before = Document(file_before)
    doc_after = Document(file_after)

    if scope_config:
        scope = resolve_scope(doc_before.paragraphs, **scope_config)
    else:
        scope = resolve_scope(doc_before.paragraphs)

    result = check_scope_integrity(
        doc_before, doc_after, scope,
        scope_description=str(scope_config or "full document")
    )

    return {
        'scope_intact': result.scope_intact,
        'requested_scope': result.requested_scope,
        'unchanged_regions': result.unchanged_regions,
        'unexpected_changes': result.unexpected_changes,
        'warnings': result.warnings,
    }


# ═══════════════════════════════════════════════════════════════════════════
# Tool 5: writing_word_format_tables
# ═══════════════════════════════════════════════════════════════════════════

def writing_word_format_tables(file_path, output_path=None, baseline_path=None, include_unknown=False):
    """Safely convert scholarly data tables to three-line format.

    Layout/figure-container tables are skipped.  If a baseline manuscript is
    supplied, border weights are inherited from it when recognizable.
    """
    if not os.path.exists(file_path):
        return {'error': f"File not found: {file_path}"}
    if baseline_path and not os.path.exists(baseline_path):
        return {'error': f"Baseline file not found: {baseline_path}"}

    from docx import Document
    from word_guard.table_format import convert_tables_to_three_line

    doc = Document(file_path)
    results = convert_tables_to_three_line(doc, baseline_path, include_unknown=include_unknown)
    save_path = output_path or file_path
    doc.save(save_path)
    validation = validate_docx_package(save_path).to_dict()

    return {
        'success': validation['valid'],
        'output': save_path,
        'tables_formatted': sum(1 for r in results if r.get('action') == 'formatted'),
        'tables_skipped': sum(1 for r in results if r.get('action') == 'skipped'),
        'baseline': baseline_path,
        'package_validation': validation,
        'details': results,
    }


def writing_word_package_validate(file_path):
    if not os.path.exists(file_path):
        return {'error': f"File not found: {file_path}"}
    return validate_docx_package(file_path).to_dict()


def writing_word_fingerprint(file_path):
    if not os.path.exists(file_path):
        return {'error': f"File not found: {file_path}"}
    return build_baseline_fingerprint(file_path).to_dict()


def writing_word_equation_audit(file_path, baseline_path=None):
    if not os.path.exists(file_path):
        return {'error': f"File not found: {file_path}"}
    if baseline_path and not os.path.exists(baseline_path):
        return {'error': f"Baseline file not found: {baseline_path}"}
    return audit_equations(file_path, baseline_path)


# ═══════════════════════════════════════════════════════════════════════════
# CLI Interface
# ═══════════════════════════════════════════════════════════════════════════

def main():
    """CLI entry point."""
    if len(sys.argv) < 2:
        print("Usage: python -m word_guard.cli <command> [args]")
        print("")
        print("Commands:")
        print("  scan <file.docx>                          - Structural scan")
        print("  audit <file.docx>                         - Writing audit")
        print("  edit <file.docx> <replacements.json>      - Edit document")
        print("  scope-check <before.docx> <after.docx>    - Check scope integrity")
        print("  format-tables <file.docx> [output.docx] [baseline.docx] - Safe three-line tables")
        print("  package-validate <file.docx>              - Validate OOXML package")
        print("  fingerprint <file.docx>                   - Baseline formatting fingerprint")
        print("  equation-audit <file.docx> [baseline.docx]- Audit native equations")
        sys.exit(1)

    command = sys.argv[1]

    if command == 'scan':
        if len(sys.argv) < 3:
            print("Usage: scan <file.docx>")
            sys.exit(1)
        result = writing_word_scan(sys.argv[2])
        print(json.dumps(result, indent=2, ensure_ascii=False))

    elif command == 'audit':
        if len(sys.argv) < 3:
            print("Usage: audit <file.docx>")
            sys.exit(1)
        result = writing_word_audit(sys.argv[2])
        print(json.dumps(result, indent=2, ensure_ascii=False))

    elif command == 'edit':
        if len(sys.argv) < 4:
            print("Usage: edit <file.docx> <replacements.json> [output.docx] [mode] [scope.json]")
            sys.exit(1)
        with open(sys.argv[3], 'r', encoding='utf-8') as f:
            replacements = json.load(f)
        output_file = sys.argv[4] if len(sys.argv) > 4 else None
        mode = sys.argv[5] if len(sys.argv) > 5 else 'text_only'
        scope_config = None
        if len(sys.argv) > 6 and sys.argv[6] not in ('null', ''):
            raw = json.loads(sys.argv[6])
            # TypeScript schema uses camelCase; Python scope resolver uses snake_case.
            keymap = {
                'startHeading': 'start_heading', 'endHeading': 'end_heading', 'heading': 'heading',
                'startParagraph': 'start_index', 'endParagraph': 'end_index',
            }
            scope_config = {keymap.get(k, k): v for k, v in raw.items()}
        result = writing_word_edit(sys.argv[2], replacements, scope_config, mode, output_file)
        print(json.dumps(result, indent=2, ensure_ascii=False))

    elif command == 'scope-check':
        if len(sys.argv) < 4:
            print("Usage: scope-check <before.docx> <after.docx> [scope.json]")
            sys.exit(1)
        scope_config = None
        if len(sys.argv) > 4 and sys.argv[4] not in ('null', ''):
            raw = json.loads(sys.argv[4])
            keymap = {'startHeading':'start_heading','endHeading':'end_heading','heading':'heading',
                      'startParagraph':'start_index','endParagraph':'end_index'}
            scope_config = {keymap.get(k, k): v for k, v in raw.items()}
        result = writing_word_scope_check(sys.argv[2], sys.argv[3], scope_config)
        print(json.dumps(result, indent=2, ensure_ascii=False))

    elif command == 'format-tables':
        if len(sys.argv) < 3:
            print("Usage: format-tables <file.docx> [output.docx] [baseline.docx]")
            sys.exit(1)
        input_file = sys.argv[2]
        output_file = sys.argv[3] if len(sys.argv) > 3 else input_file
        baseline_file = sys.argv[4] if len(sys.argv) > 4 and sys.argv[4] not in ('null','') else None
        result = writing_word_format_tables(input_file, output_file, baseline_file)
        print(json.dumps(result, indent=2, ensure_ascii=False))

    elif command == 'package-validate':
        result = writing_word_package_validate(sys.argv[2])
        print(json.dumps(result, indent=2, ensure_ascii=False))

    elif command == 'fingerprint':
        result = writing_word_fingerprint(sys.argv[2])
        print(json.dumps(result, indent=2, ensure_ascii=False))

    elif command == 'equation-audit':
        baseline_file = sys.argv[3] if len(sys.argv) > 3 else None
        result = writing_word_equation_audit(sys.argv[2], baseline_file)
        print(json.dumps(result, indent=2, ensure_ascii=False))

    else:
        print(f"Unknown command: {command}")
        sys.exit(1)


if __name__ == '__main__':
    main()
