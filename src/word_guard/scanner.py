# -*- coding: utf-8 -*-
"""
DOCX Structural Scanner.

Scans a Word document and returns structural information:
- Paragraph-level info (text, style, heading level, complex objects)
- Heading hierarchy
- Tables, figures, equations
- Section properties
- Document profile detection (manuscript/rebuttal)
"""

import re
import os
from dataclasses import dataclass, field
from typing import Optional
from docx import Document
from docx.oxml.ns import qn


@dataclass
class WordParagraphInfo:
    """Structural information for a single paragraph."""
    index: int
    text: str
    style_id: Optional[str] = None
    style_name: Optional[str] = None
    outline_level: Optional[int] = None
    is_heading: bool = False
    heading_level: Optional[int] = None
    has_equation: bool = False
    has_drawing: bool = False
    has_field: bool = False
    has_hyperlink: bool = False
    has_bookmark: bool = False
    has_footnote: bool = False
    has_comment: bool = False
    is_complex: bool = False  # True if any of the above are True
    run_count: int = 0
    char_count: int = 0
    # Heading detection signals
    is_numbered_heading: bool = False
    detected_heading_number: Optional[str] = None


@dataclass
class WordTableInfo:
    """Structural information for a table."""
    index: int  # Index in document.tables
    paragraph_index: int  # Index of the paragraph preceding the table
    rows: int = 0
    cols: int = 0
    style_id: Optional[str] = None
    text_preview: str = ""


@dataclass
class WordDocumentInfo:
    """Complete structural scan of a Word document."""
    file_path: str
    file_name: str
    paragraph_count: int = 0
    paragraphs: list = field(default_factory=list)  # list[WordParagraphInfo]
    headings: list = field(default_factory=list)  # list[WordParagraphInfo]
    tables: list = field(default_factory=list)  # list[WordTableInfo]
    section_count: int = 0
    # Document profile
    detected_profile: str = "unknown"  # manuscript/rebuttal/cover_letter/unknown
    # Aggregate counts
    total_equations: int = 0
    total_drawings: int = 0
    total_fields: int = 0
    total_hyperlinks: int = 0
    total_bookmarks: int = 0
    total_footnotes: int = 0
    total_comments: int = 0
    # Warnings
    warnings: list = field(default_factory=list)


# ── Heading Detection ─────────────────────────────────────────────────────

# Common heading number patterns: "1.", "3.1", "3.1.1", "A.", "I.", etc.
HEADING_NUMBER_RE = re.compile(
    r'^(\d+(?:\.\d+)*)\s*[.\)]\s*'          # "1.", "3.1.", "3.1.1)"
    r'|^[A-Z]\s*[.\)]\s*'                     # "A.", "B)"
    r'|^(?:I{1,3}|IV|V|VI{0,3}|IX|X)\s*[.\)]\s*'  # Roman: "I.", "II."
)


def _detect_outline_level(para_element):
    """Extract outline level from paragraph properties."""
    pPr = para_element.find(qn('w:pPr'))
    if pPr is None:
        return None
    # Check outlineLvl
    outline = pPr.find(qn('w:outlineLvl'))
    if outline is not None:
        val = outline.get(qn('w:val'))
        if val is not None:
            return int(val)
    return None


def _detect_heading_level_from_style(style_id, style_name):
    """Detect heading level from style ID or name."""
    if style_id:
        # Standard heading styles: "Heading1", "Heading 1", etc.
        m = re.search(r'heading\s*(\d+)', style_id, re.IGNORECASE)
        if m:
            return int(m.group(1))
    if style_name:
        m = re.search(r'heading\s*(\d+)', style_name, re.IGNORECASE)
        if m:
            return int(m.group(1))
    return None


def _detect_heading_level(para_element, style_id, style_name):
    """
    Detect heading level using multiple signals:
    1. outline level from XML
    2. style name/ID
    3. numbering
    4. text pattern (e.g. "3.1. Something")
    5. direct formatting (large bold font)
    """
    # Signal 1: outline level
    ol = _detect_outline_level(para_element)
    if ol is not None:
        return min(ol + 1, 9)  # outline level 0 = Heading 1

    # Signal 2: style
    level = _detect_heading_level_from_style(style_id, style_name)
    if level is not None:
        return level

    # Signal 5: direct formatting (last resort - large bold font)
    pPr = para_element.find(qn('w:pPr'))
    # Don't auto-detect from direct formatting alone - too risky
    return None


def _check_complex(para_element):
    """Check if a paragraph contains complex objects."""
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
    }
    has_eq = para_element.find('.//' + qn('m:oMath')) is not None
    has_eq_para = para_element.find('.//' + qn('m:oMathPara')) is not None
    has_drawing = para_element.find('.//' + qn('w:drawing')) is not None
    has_field = (
        para_element.find('.//' + qn('w:fldChar')) is not None or
        para_element.find('.//' + qn('w:instrText')) is not None
    )
    has_hyperlink = para_element.find(qn('w:hyperlink')) is not None
    has_bookmark = (
        para_element.find('.//' + qn('w:bookmarkStart')) is not None
    )
    has_footnote = para_element.find('.//' + qn('w:footnoteReference')) is not None
    has_comment = para_element.find('.//' + qn('w:commentReference')) is not None

    return {
        'has_equation': has_eq or has_eq_para,
        'has_drawing': has_drawing,
        'has_field': has_field,
        'has_hyperlink': has_hyperlink,
        'has_bookmark': has_bookmark,
        'has_footnote': has_footnote,
        'has_comment': has_comment,
        'is_complex': any([has_eq, has_eq_para, has_drawing, has_field,
                          has_hyperlink, has_bookmark, has_footnote, has_comment]),
    }


# ── Document Profile Detection ────────────────────────────────────────────

def _detect_profile(file_path, paragraphs):
    """Detect document profile from filename and content."""
    basename = os.path.basename(file_path).lower()

    # Filename-based detection
    if any(kw in basename for kw in ['rebuttal', 'response', 'reply', '返修', '回复']):
        return 'rebuttal'
    if any(kw in basename for kw in ['cover', 'cover_letter', '投稿']):
        return 'cover_letter'
    if any(kw in basename for kw in ['review', '审稿']):
        return 'review'

    # Content-based detection (check first 20 paragraphs)
    content_snippet = ' '.join(p.text for p in paragraphs[:20]).lower()
    if 'response to reviewer' in content_snippet or 'we thank the reviewer' in content_snippet:
        return 'rebuttal'
    if 'dear editor' in content_snippet or 'dear professor' in content_snippet:
        return 'cover_letter'

    # Default to manuscript for .docx files
    if basename.endswith('.docx'):
        return 'manuscript'

    return 'unknown'


# ── Main Scanner ──────────────────────────────────────────────────────────

def scan_docx(file_path):
    """
    Perform a structural scan of a DOCX file.

    Returns a WordDocumentInfo with detailed paragraph-level,
    heading, table, and object information.

    Args:
        file_path: Path to the .docx file.

    Returns:
        WordDocumentInfo with complete structural analysis.
    """
    doc = Document(file_path)
    info = WordDocumentInfo(
        file_path=file_path,
        file_name=os.path.basename(file_path),
    )

    # ── Scan paragraphs ──
    for i, para in enumerate(doc.paragraphs):
        style_id = para.style.style_id if para.style else None
        style_name = para.style.name if para.style else None

        # Check complex objects
        complex_check = _check_complex(para._element)

        # Detect heading level
        heading_level = _detect_heading_level(para._element, style_id, style_name)
        is_heading = heading_level is not None

        # Also check outline level directly
        outline_level = _detect_outline_level(para._element)

        # Heading number detection
        is_numbered = False
        heading_number = None
        if para.text.strip():
            m = HEADING_NUMBER_RE.match(para.text.strip())
            if m:
                is_numbered = True
                for g in m.groups():
                    if g:
                        heading_number = g
                        break

        p_info = WordParagraphInfo(
            index=i,
            text=para.text,
            style_id=style_id,
            style_name=style_name,
            outline_level=outline_level,
            is_heading=is_heading,
            heading_level=heading_level,
            has_equation=complex_check['has_equation'],
            has_drawing=complex_check['has_drawing'],
            has_field=complex_check['has_field'],
            has_hyperlink=complex_check['has_hyperlink'],
            has_bookmark=complex_check['has_bookmark'],
            has_footnote=complex_check['has_footnote'],
            has_comment=complex_check['has_comment'],
            is_complex=complex_check['is_complex'],
            run_count=len(para.runs),
            char_count=len(para.text),
            is_numbered_heading=is_numbered,
            detected_heading_number=heading_number,
        )

        info.paragraphs.append(p_info)
        if is_heading:
            info.headings.append(p_info)

        # Aggregate counts
        if complex_check['has_equation']:
            info.total_equations += 1
        if complex_check['has_drawing']:
            info.total_drawings += 1
        if complex_check['has_field']:
            info.total_fields += 1
        if complex_check['has_hyperlink']:
            info.total_hyperlinks += 1
        if complex_check['has_bookmark']:
            info.total_bookmarks += 1
        if complex_check['has_footnote']:
            info.total_footnotes += 1
        if complex_check['has_comment']:
            info.total_comments += 1

    info.paragraph_count = len(info.paragraphs)

    # ── Scan tables ──
    # Note: python-docx tables are separate from paragraphs;
    # we need to find them in the XML body
    body = doc.element.body
    tbl_elements = body.findall(qn('w:tbl'))
    for t_idx, tbl_elem in enumerate(tbl_elements):
        # Find the paragraph index of the element just before this table
        tbl_parent = tbl_elem.getparent()
        tbl_pos = list(tbl_parent).index(tbl_elem)
        # Count paragraphs before this position
        para_count_before = 0
        for child in list(tbl_parent)[:tbl_pos]:
            if child.tag == qn('w:p'):
                para_count_before += 1

        rows = len(tbl_elem.findall(qn('w:tr')))
        cols_list = tbl_elem.findall('.//' + qn('w:tc'))
        # Estimate columns from first row
        first_row = tbl_elem.find(qn('w:tr'))
        cols = len(first_row.findall(qn('w:tc'))) if first_row is not None else 0

        # Get text preview from first few cells
        cell_texts = []
        for tc in cols_list[:6]:
            t = ''.join(node.text or '' for node in tc.iter() if node.tag.endswith('}t'))
            if t.strip():
                cell_texts.append(t.strip()[:30])

        t_info = WordTableInfo(
            index=t_idx,
            paragraph_index=para_count_before,
            rows=rows,
            cols=cols,
            text_preview=' | '.join(cell_texts[:4]),
        )
        info.tables.append(t_info)

    # ── Section count ──
    sectPr_elements = body.findall('.//' + qn('w:sectPr'))
    info.section_count = len(sectPr_elements) if sectPr_elements else 1

    # ── Document profile ──
    info.detected_profile = _detect_profile(file_path, doc.paragraphs)

    # ── Warnings ──
    if info.total_fields > 0:
        info.warnings.append(
            f"Document contains {info.total_fields} field codes - "
            "editing must preserve all field boundaries"
        )
    if info.total_equations > 0:
        info.warnings.append(
            f"Document contains {info.total_equations} equations - "
            "complex paragraphs need run-aware editing"
        )
    if info.total_bookmarks > 0:
        info.warnings.append(
            f"Document contains {info.total_bookmarks} bookmarks - "
            "bookmark nodes must not be modified"
        )

    return info


def print_scan_summary(info):
    """Print a human-readable summary of the scan results."""
    print(f"\n{'='*60}")
    print(f"DOCX Structural Scan: {info.file_name}")
    print(f"{'='*60}")
    print(f"Profile: {info.detected_profile}")
    print(f"Paragraphs: {info.paragraph_count}")
    print(f"Headings: {len(info.headings)}")
    print(f"Tables: {len(info.tables)}")
    print(f"Sections: {info.section_count}")
    print(f"\nObjects:")
    print(f"  Equations: {info.total_equations}")
    print(f"  Drawings/Images: {info.total_drawings}")
    print(f"  Fields: {info.total_fields}")
    print(f"  Hyperlinks: {info.total_hyperlinks}")
    print(f"  Bookmarks: {info.total_bookmarks}")
    print(f"  Footnotes: {info.total_footnotes}")
    print(f"  Comments: {info.total_comments}")

    complex_paras = [p for p in info.paragraphs if p.is_complex]
    print(f"\nComplex paragraphs: {len(complex_paras)} / {info.paragraph_count}")

    if info.headings:
        print(f"\nHeading hierarchy:")
        for h in info.headings:
            indent = "  " * ((h.heading_level or 1) - 1)
            text_preview = h.text[:60] + "..." if len(h.text) > 60 else h.text
            print(f"  {indent}[H{h.heading_level or '?'}] p{h.index}: {text_preview}")

    if info.tables:
        print(f"\nTables:")
        for t in info.tables:
            print(f"  Table {t.index}: {t.rows}x{t.cols} at paragraph ~{t.paragraph_index}")
            if t.text_preview:
                print(f"    Preview: {t.text_preview[:80]}")

    if info.warnings:
        print(f"\nWarnings:")
        for w in info.warnings:
            print(f"  ⚠ {w}")

    print(f"{'='*60}\n")
