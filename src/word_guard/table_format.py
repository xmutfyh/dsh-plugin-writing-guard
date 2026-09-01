# -*- coding: utf-8 -*-
"""
Safe academic three-line table formatting for DOCX.

Principles
----------
1. Never assume every w:tbl is a semantic data table. Word documents often
   use borderless tables for figures/layout. In ``auto`` mode those tables are
   skipped conservatively.
2. Prefer baseline-derived border widths when a reference manuscript is
   available. Plugin defaults are only a fallback.
3. Put the three visible rules on cells (header top/header bottom/last-row
   bottom), not on table-level top/bottom borders. This avoids false bottom
   rules when a long table is split across pages by Word/LibreOffice.
4. Do not silently bold headers by default: typography belongs to the
   manuscript/template unless explicitly requested.
"""

from dataclasses import dataclass
from typing import Optional
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


@dataclass(frozen=True)
class ThreeLineSpec:
    top_sz: str = "12"       # eighths of a point: 12 = 1.5 pt
    header_sz: str = "6"     # 6 = 0.75 pt
    bottom_sz: str = "12"
    color: str = "auto"
    source: str = "plugin_default"

    def to_dict(self):
        return {
            "top_pt": int(self.top_sz) / 8,
            "header_pt": int(self.header_sz) / 8,
            "bottom_pt": int(self.bottom_sz) / 8,
            "color": self.color,
            "source": self.source,
        }


def _border_sz(container, edge):
    if container is None:
        return None
    el = container.find(qn(f"w:{edge}"))
    if el is None:
        return None
    val = (el.get(qn("w:val")) or "").lower()
    if val in ("none", "nil"):
        return None
    return el.get(qn("w:sz"))


def _table_border_sz(table, edge):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders")) if tbl_pr is not None else None
    return _border_sz(borders, edge)


def _cell_border_sz(cell, edge):
    tc_pr = cell._tc.tcPr
    borders = tc_pr.find(qn("w:tcBorders")) if tc_pr is not None else None
    return _border_sz(borders, edge)


def _first_nonempty(values):
    return next((v for v in values if v is not None), None)


def infer_three_line_spec(table, fallback=None):
    """Infer three-line rule widths from a reference/baseline table."""
    fallback = fallback or ThreeLineSpec()
    if not table.rows:
        return fallback

    top = _first_nonempty(
        [_cell_border_sz(c, "top") for c in table.rows[0].cells]
        + [_table_border_sz(table, "top")]
    )
    header = _first_nonempty(
        [_cell_border_sz(c, "bottom") for c in table.rows[0].cells]
        + [_table_border_sz(table, "insideH")]
    )
    bottom = _first_nonempty(
        [_cell_border_sz(c, "bottom") for c in table.rows[-1].cells]
        + [_table_border_sz(table, "bottom")]
    )

    if not any((top, header, bottom)):
        return fallback
    return ThreeLineSpec(
        top_sz=top or fallback.top_sz,
        header_sz=header or fallback.header_sz,
        bottom_sz=bottom or fallback.bottom_sz,
        color=fallback.color,
        source="baseline",
    )


def _set_edge(borders, edge, val, sz="0", color="auto"):
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)
    borders.append(el)


def _clear_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(old)


def _add_cell_rules(cell, top=None, bottom=None, color="auto"):
    if top is None and bottom is None:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    if top is not None:
        _set_edge(borders, "top", "single", top, color)
    if bottom is not None:
        _set_edge(borders, "bottom", "single", bottom, color)
    tc_pr.append(borders)


def _disable_table_level_borders(table):
    """Disable all table-level rules; visible rules are cell-level only."""
    tbl_pr = table._tbl.tblPr
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _set_edge(borders, edge, "none", "0", "auto")
    tbl_pr.append(borders)


def classify_table(table, preceding_text=""):
    """
    Conservatively classify a table as ``data``, ``layout`` or ``unknown``.

    This intentionally favours skipping over destructive formatting.
    ``force`` mode is available when the caller has already identified the
    target table semantically.
    """
    text = " ".join(
        p.text.strip()
        for row in table.rows
        for cell in row.cells
        for p in cell.paragraphs
        if p.text.strip()
    )
    drawing_count = len(table._tbl.findall(".//" + qn("w:drawing")))
    lower_prev = (preceding_text or "").strip().lower()
    style_name = (getattr(getattr(table, "style", None), "name", "") or "").lower()

    # Figure/layout tables commonly contain drawings and little/no data text.
    if drawing_count and len(text) < 160:
        return "layout", "contains drawing(s) with little tabular text"
    if lower_prev.startswith(("figure ", "fig. ", "fig ")):
        return "layout", "preceded by a figure caption/label"
    if "layout" in style_name:
        return "layout", "table style indicates layout"

    # Strong data-table signals.
    if lower_prev.startswith("table "):
        return "data", "preceded by a Table caption/label"
    if len(table.rows) >= 2 and len(table.columns) >= 2:
        nonempty = sum(
            1 for row in table.rows for cell in row.cells if cell.text.strip()
        )
        total = max(1, len(table.rows) * len(table.columns))
        if drawing_count == 0 and nonempty / total >= 0.50:
            return "data", "dense textual/numeric grid without drawings"

    return "unknown", "no reliable semantic signal"


def make_three_line_table(table, spec=None, bold_header=False):
    """Convert one already-selected semantic table to three-line format."""
    spec = spec or ThreeLineSpec()
    if not table.rows:
        return {"rows": 0, "cols": 0, "spec": spec.to_dict(), "header_bold_runs": 0}

    _disable_table_level_borders(table)
    for row in table.rows:
        for cell in row.cells:
            _clear_cell_borders(cell)

    # Explicit cell rules avoid page-break artefacts.
    for cell in table.rows[0].cells:
        _add_cell_rules(cell, top=spec.top_sz, bottom=spec.header_sz, color=spec.color)
    for cell in table.rows[-1].cells:
        # A one-row table needs top/header/bottom on the same cell. Rebuild it.
        if len(table.rows) == 1:
            _clear_cell_borders(cell)
            _add_cell_rules(cell, top=spec.top_sz, bottom=spec.bottom_sz, color=spec.color)
        else:
            _add_cell_rules(cell, bottom=spec.bottom_sz, color=spec.color)

    header_bold_count = 0
    if bold_header:
        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if not run.bold:
                        run.bold = True
                        header_bold_count += 1

    return {
        "rows": len(table.rows),
        "cols": len(table.columns),
        "spec": spec.to_dict(),
        "header_bold_runs": header_bold_count,
    }


def convert_all_tables_to_three_line(
    doc, baseline_doc=None, selection="auto", bold_header=False
):
    """
    Format semantic data tables.

    selection:
      - ``auto``  : classify and skip layout/unknown tables (safe default)
      - ``all``   : force every table (legacy/explicit behaviour)

    If ``baseline_doc`` is provided, corresponding baseline table border widths
    are inherited when detectable.
    """
    results = []
    preceding = ""
    # python-docx table indices are stable for doc.tables; use XML order only
    # to find a nearby preceding paragraph caption.
    body_children = list(doc._element.body)
    table_to_prev = {}
    prev_text = ""
    table_i = 0
    for child in body_children:
        if child.tag == qn("w:p"):
            texts = child.findall(".//" + qn("w:t"))
            candidate = "".join(t.text or "" for t in texts).strip()
            if candidate:
                prev_text = candidate
        elif child.tag == qn("w:tbl"):
            table_to_prev[table_i] = prev_text
            table_i += 1

    for i, table in enumerate(doc.tables):
        kind, reason = classify_table(table, table_to_prev.get(i, ""))
        if selection != "all" and kind != "data":
            results.append({
                "index": i, "action": "skipped", "classification": kind,
                "reason": reason,
            })
            continue

        spec = ThreeLineSpec()
        if baseline_doc is not None and i < len(baseline_doc.tables):
            spec = infer_three_line_spec(baseline_doc.tables[i], fallback=spec)

        r = make_three_line_table(table, spec=spec, bold_header=bold_header)
        r.update({
            "index": i, "action": "formatted",
            "classification": kind, "reason": reason,
        })
        results.append(r)
    return results
