# -*- coding: utf-8 -*-
"""Baseline formatting fingerprint for scholarly DOCX revision.

The fingerprint is intentionally descriptive, not prescriptive: when a baseline
manuscript exists, its formatting is authoritative.  This module records the
relevant Word/OOXML settings so later formatters can inherit rather than invent.
"""
from __future__ import annotations

from dataclasses import asdict, dataclass, field
from hashlib import sha256
import zipfile
from lxml import etree

from docx import Document
from docx.oxml.ns import qn


@dataclass
class BaselineFingerprint:
    file: str
    styles_sha256: str | None = None
    settings_sha256: str | None = None
    numbering_sha256: str | None = None
    math_font: str | None = None
    default_paragraph_style: str | None = None
    section_geometry: list = field(default_factory=list)
    table_signatures: list = field(default_factory=list)
    equation_paragraph_styles: list = field(default_factory=list)

    def to_dict(self):
        return asdict(self)


def _hash_part(zf, name):
    try:
        return sha256(zf.read(name)).hexdigest()
    except KeyError:
        return None


def _edge_spec(container, edge):
    if container is None:
        return None
    el = container.find(qn(f'w:{edge}'))
    if el is None:
        return None
    return {
        'val': el.get(qn('w:val')),
        'sz': el.get(qn('w:sz')),
        'color': el.get(qn('w:color')),
    }


def _table_signature(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    sig = {
        'rows': len(table.rows),
        'cols': len(table.columns),
        'style': table.style.style_id if table.style else None,
        'table_borders': {e: _edge_spec(tblBorders, e) for e in ('top','bottom','left','right','insideH','insideV')},
        'first_row_cell_borders': [],
        'last_row_cell_borders': [],
    }
    if table.rows:
        for cell in table.rows[0].cells:
            tcBorders = cell._tc.get_or_add_tcPr().find(qn('w:tcBorders'))
            sig['first_row_cell_borders'].append({e: _edge_spec(tcBorders, e) for e in ('top','bottom')})
        for cell in table.rows[-1].cells:
            tcBorders = cell._tc.get_or_add_tcPr().find(qn('w:tcBorders'))
            sig['last_row_cell_borders'].append({e: _edge_spec(tcBorders, e) for e in ('bottom',)})
    return sig


def build_baseline_fingerprint(path: str) -> BaselineFingerprint:
    doc = Document(path)
    fp = BaselineFingerprint(file=path)
    with zipfile.ZipFile(path) as zf:
        fp.styles_sha256 = _hash_part(zf, 'word/styles.xml')
        fp.settings_sha256 = _hash_part(zf, 'word/settings.xml')
        fp.numbering_sha256 = _hash_part(zf, 'word/numbering.xml')
        try:
            root = etree.fromstring(zf.read('word/settings.xml'))
            mathPr = root.find('.//' + qn('m:mathPr'))
            if mathPr is not None:
                mathFont = mathPr.find(qn('m:mathFont'))
                if mathFont is not None:
                    fp.math_font = mathFont.get(qn('m:val'))
        except Exception:
            pass

    try:
        fp.default_paragraph_style = doc.styles['Normal'].style_id
    except Exception:
        pass

    for sec in doc.sections:
        fp.section_geometry.append({
            'page_width': sec.page_width,
            'page_height': sec.page_height,
            'left_margin': sec.left_margin,
            'right_margin': sec.right_margin,
            'top_margin': sec.top_margin,
            'bottom_margin': sec.bottom_margin,
        })

    fp.table_signatures = [_table_signature(t) for t in doc.tables]

    eq_styles = []
    for p in doc.paragraphs:
        if p._element.find('.//' + qn('m:oMath')) is not None or p._element.find('.//' + qn('m:oMathPara')) is not None:
            sid = p.style.style_id if p.style else None
            if sid not in eq_styles:
                eq_styles.append(sid)
    fp.equation_paragraph_styles = eq_styles
    return fp
